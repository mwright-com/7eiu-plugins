#!/usr/bin/env python3
"""
Render a 7EIU assessment into the benchmark's layout.

    python3 build_assessment.py assessment.md -o output/

Input is the markdown produced by the orchestrator, with the YAML front matter block
described in references/output_templates.md. Output is a .docx. Convert to PDF with:

    soffice --headless --convert-to pdf --outdir output output/<name>.docx

Diagrams: any ![alt](figures/x.png) line becomes a centered figure with the alt text as
its caption. Render the HTML diagrams in references/ to PNG first, with

    python3 build_assessment.py --figures assessment.md -o output/

which needs a headless chromium on PATH (chromium, chromium-browser, or
google-chrome). Without one, the figure lines are skipped and a note is printed, so the
document still builds.

Requires: python-docx.
"""

import argparse
import os
import re
import shutil
import subprocess
import sys
import tempfile

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    sys.exit("python-docx is required:  pip install python-docx")


# ---------------------------------------------------------------- layout constants

PAGE_W, PAGE_H = Inches(8.5), Inches(11.0)
MARGIN = Inches(1.0)
HEADER_MARGIN = Inches(0.5)

BODY_FONT = "Arial"
BODY_SIZE = Pt(11)
LINE_SPACING = 1.5

BLUE_HEAD = RGBColor(0x1F, 0x5F, 0xBF)
BLUE_TITLE = RGBColor(0x0B, 0x3D, 0x91)
GREY = RGBColor(0x6E, 0x75, 0x7D)
INK = RGBColor(0x11, 0x14, 0x17)

SZ = {
    "header": Pt(9),
    "kicker": Pt(13),
    "title": Pt(28),
    "subtitle": Pt(13),
    "block_label": Pt(11),
    "principle": Pt(10.5),
    "caption": Pt(9),
}

REF_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "references")


# ---------------------------------------------------------------- front matter

def parse_front_matter(text):
    if not text.startswith("---"):
        return {}, text
    end = text.find("\n---", 3)
    if end == -1:
        return {}, text
    raw, body = text[3:end], text[end + 4:]
    meta = {}
    stack = [(-1, meta)]
    for line in raw.splitlines():
        if not line.strip() or line.strip().startswith("#"):
            continue
        indent = len(line) - len(line.lstrip())
        s = line.strip()
        while stack and indent <= stack[-1][0]:
            stack.pop()
        parent = stack[-1][1]
        if s.startswith("- "):
            parent.setdefault("_list", []).append(s[2:].strip().strip('"'))
            continue
        if ":" not in s:
            continue
        k, _, v = s.partition(":")
        k, v = k.strip(), v.strip().strip('"')
        if v:
            parent[k] = v
        else:
            child = {}
            parent[k] = child
            stack.append((indent, child))
    return _flatten_lists(meta), body.lstrip("\n")


def _flatten_lists(d):
    if isinstance(d, dict):
        if set(d.keys()) == {"_list"}:
            return d["_list"]
        return {k: _flatten_lists(v) for k, v in d.items()}
    return d


# ---------------------------------------------------------------- docx helpers

def set_base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = BODY_FONT
    st.font.size = BODY_SIZE
    st.font.color.rgb = INK
    st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    pf = st.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(10)


def para(doc, text="", size=None, bold=False, italic=False, color=None,
         align=None, space_before=None, space_after=None, spacing=None,
         tracking=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    if space_before is not None:
        pf.space_before = space_before
    if space_after is not None:
        pf.space_after = space_after
    pf.line_spacing = LINE_SPACING if spacing is None else spacing
    if text:
        add_runs(p, text, size=size, bold=bold, italic=italic, color=color,
                 tracking=tracking)
    return p


INLINE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|\[[^\]]+\]\([^)]+\))")


def add_runs(p, text, size=None, bold=False, italic=False, color=None, tracking=None):
    """Render **bold**, *italic* and [label](url) inside a paragraph."""
    for chunk in INLINE.split(text):
        if not chunk:
            continue
        b, i, label = bold, italic, chunk
        link = None
        if chunk.startswith("**") and chunk.endswith("**"):
            b, label = True, chunk[2:-2]
        elif chunk.startswith("*") and chunk.endswith("*"):
            i, label = True, chunk[1:-1]
        else:
            m = re.fullmatch(r"\[([^\]]+)\]\(([^)]+)\)", chunk)
            if m:
                label, link = m.group(1), m.group(2)
        r = p.add_run(label)
        r.bold, r.italic = b, i
        if size:
            r.font.size = size
        if color is not None:
            r.font.color.rgb = color
        if link:
            r.font.color.rgb = BLUE_HEAD
            r.font.underline = True
        if tracking:
            sp = OxmlElement("w:spacing")
            sp.set(qn("w:val"), str(int(tracking * 20)))
            r._element.get_or_add_rPr().append(sp)


def hrule(doc, color="1F5FBF", size=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:color"), color)
    bottom.set(qn("w:space"), "1")
    pbdr.append(bottom)
    p._p.get_or_add_pPr().append(pbdr)
    return p


def build_header(section, meta, mark=None):
    by = meta.get("prepared_by", {})
    # A borderless two-cell table, not tabs. The Header style carries inherited tab
    # stops that a paragraph-level override does not reliably beat, and LibreOffice
    # honours the inherited ones, so the right-hand block ends up mid-page.
    section.header.paragraphs[0].text = ""
    t = section.header.add_table(rows=1, cols=2, width=PAGE_W - 2 * MARGIN)
    t.autofit = False
    cells = t.rows[0].cells
    cells[0].width = Inches(3.6)
    cells[1].width = Inches(2.9)

    pairs = (
        (cells[0], "%s  - %s" % (by.get("name", ""), by.get("city", "")),
         WD_ALIGN_PARAGRAPH.LEFT),
        (cells[1], "%s - %s" % (by.get("email", ""), by.get("phone", "")),
         WD_ALIGN_PARAGRAPH.RIGHT),
    )
    for cell, text, align in pairs:
        cp = cell.paragraphs[0]
        cp.alignment = align
        cp.paragraph_format.line_spacing = 1.0
        cp.paragraph_format.space_after = Pt(0)
        r = cp.add_run(text)
        r.font.size = SZ["header"]
        r.font.color.rgb = GREY

    p = section.header.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

    if mark and os.path.exists(mark):
        m = section.header.add_paragraph()
        m.alignment = WD_ALIGN_PARAGRAPH.CENTER
        m.paragraph_format.space_before = Pt(4)
        m.paragraph_format.space_after = Pt(0)
        m.add_run().add_picture(mark, height=Inches(0.42))


def find_mark(md_path):
    """Brand mark, centered under the header on every page.

    A mark beside the assessment wins, so a white-label engagement can drop in the
    client's own. Otherwise the packaged MWRIGHT INC mark is used.
    """
    here = os.path.dirname(os.path.abspath(md_path))
    for cand in ("mark.png", "assets/mark.png", "figures/mark.png",
                 "logo.png", "assets/logo.png"):
        path = os.path.join(here, cand)
        if os.path.exists(path):
            return path
    packaged = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                            "..", "assets", "mark.png")
    return packaged if os.path.exists(packaged) else None


def add_page_numbers(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    for el, attrs, txt in (
        ("w:fldChar", {"w:fldCharType": "begin"}, None),
        ("w:instrText", {"xml:space": "preserve"}, "PAGE"),
        ("w:fldChar", {"w:fldCharType": "end"}, None),
    ):
        e = OxmlElement(el)
        for k, v in attrs.items():
            e.set(qn(k), v)
        if txt:
            e.text = txt
        r._r.append(e)
    r.font.size = Pt(10)


# ---------------------------------------------------------------- cover

def render_cover(doc, meta):
    para(doc, space_after=Pt(120))
    if meta.get("kicker"):
        para(doc, meta["kicker"], size=SZ["kicker"], bold=True, color=GREY,
             align=WD_ALIGN_PARAGRAPH.CENTER, tracking=3, space_after=Pt(14))
    title = meta.get("client", "")
    if meta.get("engagement"):
        title += " - " + meta["engagement"]
    para(doc, title, size=SZ["title"], bold=True, color=BLUE_TITLE,
         align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0, space_after=Pt(18))
    if meta.get("subtitle"):
        para(doc, meta["subtitle"], size=SZ["subtitle"], color=GREY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(200))

    hrule(doc)
    t = doc.add_table(rows=1, cols=2)
    t.autofit = True
    for cell, key, label in ((t.rows[0].cells[0], "prepared_for", "PREPARED FOR"),
                             (t.rows[0].cells[1], "prepared_by", "PREPARED BY")):
        blk = meta.get(key, {})
        p0 = cell.paragraphs[0]
        p0.paragraph_format.line_spacing = 1.0
        add_runs(p0, label, size=SZ["block_label"], color=BLUE_HEAD, tracking=1)
        lines = [("**%s**" % blk.get("name", ""))]
        if blk.get("title"):
            lines.append(blk["title"])
        if blk.get("company"):
            lines.append("%s %s %s" % (blk["company"], chr(0x00B7), blk.get("city", "")))
        for a in blk.get("address", []) or []:
            lines.append(a)
        for k in ("email", "phone"):
            if blk.get(k):
                lines.append(blk[k])
        for ln in lines:
            q = cell.add_paragraph()
            q.paragraph_format.line_spacing = 1.0
            q.paragraph_format.space_after = Pt(0)
            add_runs(q, ln)
    doc.add_page_break()


# ---------------------------------------------------------------- body

FIG = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")
PRINCIPLE = re.compile(r"^\*\*(PRINCIPLE\s+\d+:[^*]+)\*\*\s*$")


def render_body(doc, body, base_dir, missing):
    for block in re.split(r"\n\s*\n", body):
        block = block.strip()
        if not block:
            continue

        m = FIG.match(block)
        if m:
            alt, src = m.group(1), m.group(2)
            path = src if os.path.isabs(src) else os.path.join(base_dir, src)
            if os.path.exists(path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(path, width=PAGE_W - 2 * MARGIN)
                para(doc, alt, size=SZ["caption"], color=GREY,
                     align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0, space_after=Pt(14))
            else:
                missing.append(src)
            continue

        m = PRINCIPLE.match(block)
        if m:
            para(doc, m.group(1), size=SZ["principle"], bold=True, color=BLUE_HEAD,
                 space_before=Pt(22), space_after=Pt(10), spacing=1.0)
            continue

        if block.startswith("### "):
            para(doc, block[4:], size=Pt(14), bold=True, color=BLUE_TITLE,
                 space_before=Pt(16), space_after=Pt(6), spacing=1.0)
            continue
        if block.startswith("## "):
            para(doc, block[3:], size=Pt(16), bold=True, color=BLUE_TITLE,
                 space_before=Pt(16), space_after=Pt(8), spacing=1.0)
            continue

        lines = block.splitlines()
        if all(ln.strip().startswith(("- ", "* ")) for ln in lines):
            for ln in lines:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.line_spacing = LINE_SPACING
                p.paragraph_format.space_after = Pt(2)
                add_runs(p, ln.strip()[2:], italic=True)
            continue

        para(doc, " ".join(ln.strip() for ln in lines))


# ---------------------------------------------------------------- figures

def find_browser():
    for name in ("chromium", "chromium-browser", "google-chrome", "google-chrome-stable"):
        path = shutil.which(name)
        if path:
            return path
    # Playwright's bundled chromium, if this environment has one.
    import glob as _glob
    for pat in ("/opt/pw-browsers/chromium-*/chrome-linux/chrome",
                os.path.expanduser("~/.cache/ms-playwright/chromium-*/chrome-linux/chrome"),
                "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"):
        hits = sorted(_glob.glob(pat))
        if hits:
            return hits[-1]
    return None


def trim_png(path, pad=18):
    """Crop the blank margin a fixed-size screenshot leaves under short content."""
    try:
        from PIL import Image, ImageChops
    except ImportError:
        return
    im = Image.open(path).convert("RGB")
    bg = Image.new("RGB", im.size, im.getpixel((0, 0)))
    box = ImageChops.difference(im, bg).getbbox()
    if not box:
        return
    l, t, r, b = box
    l, t = max(0, l - pad), max(0, t - pad)
    r, b = min(im.width, r + pad), min(im.height, b + pad)
    im.crop((l, t, r, b)).save(path)


def render_figures(md_path, out_dir, meta, sub):
    """Render the reference HTML diagrams to PNG next to the markdown."""
    browser = find_browser()
    fig_dir = os.path.join(os.path.dirname(os.path.abspath(md_path)), "figures")
    os.makedirs(fig_dir, exist_ok=True)
    if not browser:
        print("no headless browser found; skipping figure rendering", file=sys.stderr)
        return

    jobs = [("diagram_seven_principles.html", "seven_principles.png", 1300, 760),
            ("diagram_marketing_garden.html", "marketing_garden.png", 1000, 1180)]
    for src, out, w, h in jobs:
        tpl_path = os.path.join(REF_DIR, src)
        if not os.path.exists(tpl_path):
            continue
        with open(tpl_path, encoding="utf-8") as fh:
            html = fh.read()
        for k, v in sub.items():
            html = html.replace("{{%s}}" % k, str(v))
        if "{{" in html:
            # Placeholders left over: this document does not include that figure.
            continue
        with tempfile.NamedTemporaryFile("w", suffix=".html", delete=False,
                                         encoding="utf-8") as tmp:
            tmp.write(html)
            tmp_path = tmp.name
        dest = os.path.join(fig_dir, out)
        subprocess.run(
            [browser, "--headless", "--disable-gpu", "--no-sandbox", "--hide-scrollbars",
             "--force-device-scale-factor=2",
             "--screenshot=" + dest, "--window-size=%d,%d" % (w, h),
             "file://" + tmp_path],
            check=False, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        os.unlink(tmp_path)
        if os.path.exists(dest):
            trim_png(dest)
            print("figure:", dest)


# ---------------------------------------------------------------- main

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("markdown")
    ap.add_argument("-o", "--outdir", default="output")
    ap.add_argument("--figures", action="store_true",
                    help="render the reference diagrams to PNG before building")
    args = ap.parse_args()

    with open(args.markdown, encoding="utf-8") as fh:
        meta, body = parse_front_matter(fh.read())

    os.makedirs(args.outdir, exist_ok=True)

    if args.figures:
        sub = {"CLIENT": meta.get("client", "").upper(),
               "FOCUS": meta.get("focus_law", "7")}
        # Optional sidecar: garden.json beside the markdown fills the marketing-garden
        # figure. Keys: anchor, law, then six {asset, use} entries in slot order
        # (PROOF, CAPTURE, DIAGNOSTIC, AUTHORITY, CADENCE, ACCESS), plus foot.
        garden = os.path.join(os.path.dirname(os.path.abspath(args.markdown)),
                              "garden.json")
        if os.path.exists(garden):
            import json
            with open(garden, encoding="utf-8") as fh:
                g = json.load(fh)
            sub["ANCHOR"] = g.get("anchor", "")
            sub["LAW"] = g.get("law", "")
            sub["FOOT"] = g.get("foot", "")
            for i, slot in enumerate(g.get("slots", [])[:6], start=1):
                sub["A%d" % i] = slot.get("asset", "")
                sub["D%d" % i] = slot.get("use", "")
        render_figures(args.markdown, args.outdir, meta, sub)

    doc = Document()
    set_base_style(doc)
    s = doc.sections[0]
    s.page_width, s.page_height = PAGE_W, PAGE_H
    s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = MARGIN
    s.header_distance = HEADER_MARGIN
    build_header(s, meta, find_mark(args.markdown))
    add_page_numbers(s)

    quick = meta.get("doc_type", "full") == "quick"
    if not quick:
        render_cover(doc, meta)

    missing = []
    render_body(doc, body, os.path.dirname(os.path.abspath(args.markdown)), missing)

    slug = re.sub(r"[^A-Za-z0-9]+", "_",
                  "%s %s" % (meta.get("client", "Assessment"),
                             "Quick Read" if quick else "Assessment")).strip("_")
    out = os.path.join(args.outdir, slug + ".docx")
    doc.save(out)
    print("wrote", out)
    for src in missing:
        print("missing figure, skipped:", src, file=sys.stderr)
    print("pdf:  soffice --headless --convert-to pdf --outdir %s %s" % (args.outdir, out))


if __name__ == "__main__":
    main()
